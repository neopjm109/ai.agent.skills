#!/usr/bin/env python3
"""Extract paragraphs (with heading styles) and tables from a .docx document.

Prints each non-empty paragraph as `[StyleName] text` and each table as a
`--- TABLE n ---` block of pipe-joined rows. Prefers python-docx; falls back to
parsing word/document.xml (a .docx is a ZIP) with the standard library only, so
the sole requirement is python3.

Usage: python3 scripts/extract.py <file.docx>
"""
import sys


def with_python_docx(path):
    from docx import Document
    doc = Document(path)
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"[{p.style.name}] {p.text}")
    for i, t in enumerate(doc.tables):
        print(f"--- TABLE {i} ---")
        for row in t.rows:
            print(" | ".join(c.text for c in row.cells))


def with_stdlib(path):
    import zipfile
    from xml.etree import ElementTree as ET

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def para_text(p):
        return "".join(t.text or "" for t in p.iter(f"{W}t"))

    with zipfile.ZipFile(path) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    body = root.find(f"{W}body")
    if body is None:
        return
    table_no = 0
    for el in body:
        if el.tag == f"{W}p":
            text = para_text(el).strip()
            if not text:
                continue
            style = el.find(f"{W}pPr/{W}pStyle")
            name = style.get(f"{W}val") if style is not None else "Normal"
            print(f"[{name}] {text}")
        elif el.tag == f"{W}tbl":
            print(f"--- TABLE {table_no} ---")
            table_no += 1
            for tr in el.findall(f"{W}tr"):
                cells = ["".join(para_text(p) for p in tc.findall(f"{W}p"))
                         for tc in tr.findall(f"{W}tc")]
                print(" | ".join(cells))


def main():
    if len(sys.argv) != 2:
        sys.exit("usage: extract.py <file.docx>")
    path = sys.argv[1]
    try:
        with_python_docx(path)
    except ImportError:
        with_stdlib(path)


if __name__ == "__main__":
    main()
